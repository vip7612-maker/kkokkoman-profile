"""
강사 김춘흥 강사 관리카드 — Word(.docx) 양식.
샘플(김혜정 교육청 강사 관리카드) 표 구조를 그대로 재현.
- 주민번호/계좌/휴대폰/주소: 환경변수 주입 (커밋 금지)
- 동의 날짜는 비움
- '강원특별자치도홍천교육지원청교육장 귀하' 라인 제외
- 동의서 '서명 또는 인' 자리에 도장 이미지(stamp.png) 삽입
"""
import os
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/kj.lee/dev/꼬꼬맨 프로필페이지/app/scripts/proposal-output/instructor-card.docx"
STAMP = "/Users/kj.lee/dev/꼬꼬맨 프로필페이지/app/scripts/proposal-output/stamp.png"

# ── 강사 정보 ──
NAME = "김춘흥"
RRN = os.environ.get("INSTRUCTOR_RRN", "")
ACCOUNT = os.environ.get("INSTRUCTOR_ACCOUNT", "")
PHONE = os.environ.get("INSTRUCTOR_PHONE", "")
HOME_ADDR = os.environ.get("INSTRUCTOR_HOME_ADDR", "")
FIELD = "팬플룻 연주, 음악심리"
HOMEPAGE = "https://kkokkoman.vercel.app"
EMAIL = "jaru1125@naver.com"
MAJOR = ""

# 한글 폰트 (Word/Pages 모두에서 표준)
KOR_FONT = "맑은 고딕"
HEADER_BG = "DCEAF1"  # 연한 청록 (샘플과 유사)


def set_cell_bg(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_vertical_align(cell, align="center"):
    tc_pr = cell._tc.get_or_add_tcPr()
    v = OxmlElement("w:vAlign")
    v.set(qn("w:val"), align)
    tc_pr.append(v)


def set_cell_text_direction_vertical(cell):
    """세로쓰기 (TBLR 또는 TBVERT)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    td = OxmlElement("w:textDirection")
    td.set(qn("w:val"), "tbRl")
    tc_pr.append(td)


def set_run_font(run, size_pt=10, bold=False, color=(0, 0, 0)):
    run.font.name = KOR_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), KOR_FONT)
    rfonts.set(qn("w:hAnsi"), KOR_FONT)
    rfonts.set(qn("w:eastAsia"), KOR_FONT)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)


def cell_text(cell, text, *, size=10, bold=False, align="center", header=False, color=(0, 0, 0)):
    cell.text = ""  # 기존 단락 정리
    p = cell.paragraphs[0]
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text or "")
    set_run_font(run, size_pt=size, bold=bold or header, color=color)
    set_cell_vertical_align(cell, "center")
    if header:
        set_cell_bg(cell, HEADER_BG)


def cell_multiline(cell, lines, *, size=9.5, header=False):
    cell.text = ""
    if header:
        set_cell_bg(cell, HEADER_BG)
    for i, ln in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(ln)
        set_run_font(run, size_pt=size)


def set_table_borders(table):
    tbl_pr = table._element.find(qn("w:tblPr"))
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tbl_pr.append(borders)


# 1cm == 567 twips
def cm_to_twips(cm: float) -> int:
    return int(cm * 567)


def set_table_width_cm(table, total_cm: float):
    """표 자체의 너비를 명시적으로 dxa(twips) 단위로 강제."""
    tbl_pr = table._element.find(qn("w:tblPr"))
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(cm_to_twips(total_cm)))
    tbl_w.set(qn("w:type"), "dxa")
    # 고정 레이아웃 (셀 너비 그대로 적용)
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_cell_width_cm(cell, cm: float):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(cm_to_twips(cm)))
    tc_w.set(qn("w:type"), "dxa")


def set_row_height_cm(row, cm: float, rule: str = "atLeast"):
    tr_pr = row._tr.get_or_add_trPr()
    tr_h = tr_pr.find(qn("w:trHeight"))
    if tr_h is None:
        tr_h = OxmlElement("w:trHeight")
        tr_pr.append(tr_h)
    tr_h.set(qn("w:val"), str(cm_to_twips(cm)))
    tr_h.set(qn("w:hRule"), rule)


def apply_widths(table, col_widths_cm: list, row_height_cm: float = 0.85):
    """표 전체와 각 셀에 너비/높이를 일관되게 적용."""
    total_cm = sum(col_widths_cm)
    set_table_width_cm(table, total_cm)
    for row in table.rows:
        set_row_height_cm(row, row_height_cm, rule="atLeast")
        for i, cw in enumerate(col_widths_cm):
            if i < len(row.cells):
                set_cell_width_cm(row.cells[i], cw)


def main():
    doc = Document()

    # 페이지 여백
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # 제목
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("강사 관리카드")
    set_run_font(title_run, size_pt=22, bold=True)
    title_run.font.underline = True
    title_p.paragraph_format.space_after = Pt(6)

    # ── 메인 정보 표 ──
    # 4열 균등: [라벨1][값1][라벨2][값2]
    info_rows = 4
    info = doc.add_table(rows=info_rows, cols=4)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = False
    # 4열 너비 합 = 17 cm (라벨 2.5 + 값 6.0)
    apply_widths(info, [2.5, 6.0, 2.5, 6.0], row_height_cm=0.8)

    pairs = [
        ("성  명", NAME, "주민등록번호", RRN),
        ("강의분야", FIELD, "계 좌 번 호", ACCOUNT),
        ("전  공", MAJOR, "휴 대 폰", PHONE),
        ("홈페이지", HOMEPAGE, "이 메 일", EMAIL),
    ]
    for r, (l1, v1, l2, v2) in enumerate(pairs):
        cell_text(info.rows[r].cells[0], l1, header=True, size=10)
        cell_text(info.rows[r].cells[1], v1, size=11, bold=(r == 0))
        cell_text(info.rows[r].cells[2], l2, header=True, size=10)
        cell_text(info.rows[r].cells[3], v2, size=11)
    set_table_borders(info)

    # ── 자택 표 (2행) — 근무처 표는 제거 ──
    home_cols = [1.5, 1.3, 4.7, 1.3, 2.2, 2.0, 4.0]
    home = doc.add_table(rows=2, cols=7)
    home.alignment = WD_TABLE_ALIGNMENT.CENTER
    home.autofit = False
    apply_widths(home, home_cols, row_height_cm=0.8)

    home_label_cell = home.rows[0].cells[0]
    home_label_cell.merge(home.rows[1].cells[0])
    cell_text(home_label_cell, "자 택", header=True, size=10)

    cell_text(home.rows[0].cells[1], "주소", header=True, size=10)
    addr_h = home.rows[0].cells[2].merge(home.rows[0].cells[3]).merge(home.rows[0].cells[4])
    cell_text(addr_h, HOME_ADDR, size=10.5, align="left")
    cell_text(home.rows[0].cells[5], "우편번호", header=True, size=10)
    cell_text(home.rows[0].cells[6], "", size=10)

    blank_h = home.rows[1].cells[1].merge(home.rows[1].cells[2]).merge(home.rows[1].cells[3]).merge(home.rows[1].cells[4])
    cell_text(blank_h, "", size=10)
    cell_text(home.rows[1].cells[5], "전화번호", header=True, size=10)
    cell_text(home.rows[1].cells[6], PHONE, size=10.5)
    set_table_borders(home)

    # ── 학력 및 전공 (헤더 1행 + 데이터 4행) ──
    edu = doc.add_table(rows=5, cols=4)
    edu.alignment = WD_TABLE_ALIGNMENT.CENTER
    edu.autofit = False
    apply_widths(edu, [2.0, 7.0, 5.5, 2.5], row_height_cm=0.7)

    # 좌측 5행 병합 ("학력 및 전공")
    edu_label_cell = edu.rows[0].cells[0]
    for r in range(1, 5):
        edu_label_cell = edu_label_cell.merge(edu.rows[r].cells[0])
    cell_text(edu_label_cell, "학 력\n및\n전 공", header=True, size=10)

    cell_text(edu.rows[0].cells[1], "학   교   명", header=True)
    cell_text(edu.rows[0].cells[2], "전 공 과 목", header=True)
    cell_text(edu.rows[0].cells[3], "학    위", header=True)

    edu_data = [
        ("세종사이버대학교", "", "졸업"),
        ("", "", ""),
        ("", "", ""),
        ("", "", ""),
    ]
    for i, (s, m, d) in enumerate(edu_data, start=1):
        cell_text(edu.rows[i].cells[1], s, size=10.5)
        cell_text(edu.rows[i].cells[2], m, size=10.5)
        cell_text(edu.rows[i].cells[3], d, size=10.5)
    set_table_borders(edu)

    # ── 주요경력 (좌측 라벨 + 우측 본문 병합) ──
    career = doc.add_table(rows=1, cols=2)
    career.alignment = WD_TABLE_ALIGNMENT.CENTER
    career.autofit = False
    apply_widths(career, [2.0, 15.0], row_height_cm=3.6)

    cell_text(career.rows[0].cells[0], "주요경력", header=True, size=10)
    cell_multiline(career.rows[0].cells[1], [
        "전국 팬플룻 대회 금상 30회 수상",
        "상지대학교 사회복지학과 음악심리 특강",
        "현대글로비스 문막지부, 산림항공본부 강연",
        "지식나눔 시브아 포럼, 강의소림 밥포럼 출강",
        "내안에병원(정신건강의학) 정기 음악 치유 시간",
        "원주교도소, 청주여자교도소 음악·인문 강연",
        "MBC 강원365 출연 외 다수",
    ], size=9.5)
    set_table_borders(career)

    # ── 동의 박스 (단일 셀) — 한 페이지 내 컴팩트 ──
    consent = doc.add_table(rows=1, cols=1)
    consent.alignment = WD_TABLE_ALIGNMENT.CENTER
    consent.autofit = False
    apply_widths(consent, [17.0], row_height_cm=0.5)
    set_table_borders(consent)
    consent_cell = consent.rows[0].cells[0]
    consent_cell.text = ""

    def add_p(cell, text, *, size=9, bold=False, align="left", indent=0,
              space_before=0, space_after=0, line_pt=11):
        p = cell.add_paragraph()
        if align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if indent:
            p.paragraph_format.left_indent = Pt(indent)
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        # 단일 행 라인 간격
        p.paragraph_format.line_spacing = Pt(line_pt)
        run = p.add_run(text)
        set_run_font(run, size_pt=size, bold=bold)
        return p

    # 셀의 첫 paragraph 컴팩트 설정 후 헤더 작성
    p0 = consent_cell.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after = Pt(1)
    p0.paragraph_format.line_spacing = Pt(12)
    set_run_font(p0.add_run("■ 개인정보 수집ㆍ이용 동의"), size_pt=10.5, bold=True)

    add_p(consent_cell, "  ◇ 개인정보 수집ㆍ이용 목적: 신청접수, 연락, 수료 관리, 강사료 지급", size=8.5)
    add_p(consent_cell, "  ◇ 수집항목: 성명, 주민등록번호, 직장명(직급/직위 포함), 주소, 연락처, 계좌번호, 경력 및 자격사항", size=8.5)
    add_p(consent_cell, "  ◇ 보유 및 이용기간: 보존기한 내 또는 정보제공자 철회 요청까지", size=8.5)
    add_p(consent_cell, "  ◇ 본 개인정보 수집에 대한 동의를 거부하실 수 있으며, 이 경우 강사 등록 및 강사료 지급이 제한될 수 있음", size=8.5)

    add_p(consent_cell, "    개인정보 수집ㆍ이용 동의      ■ 동의       □ 동의하지 않음",
          size=9.5, align="center", bold=True, space_before=2)
    add_p(consent_cell, "          년       월       일", size=9.5, align="center")

    # 성명 + 도장
    sig_p = consent_cell.add_paragraph()
    sig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_p.paragraph_format.space_before = Pt(0)
    sig_p.paragraph_format.space_after = Pt(0)
    sig_p.paragraph_format.line_spacing = Pt(14)
    run = sig_p.add_run("성명          ")
    set_run_font(run, size_pt=10)
    run2 = sig_p.add_run(NAME)
    set_run_font(run2, size_pt=10, bold=True)
    run3 = sig_p.add_run("    (서명 또는 인) ")
    set_run_font(run3, size_pt=10)
    if os.path.exists(STAMP):
        run4 = sig_p.add_run()
        run4.add_picture(STAMP, width=Cm(1.2))

    add_p(consent_cell,
          "※ 수집한 개인정보는 정보주체의 동의 없이 수집한 목적 외로 사용하거나 제3자에게 제공되지 않으며 보유 및 이용기간 만료 이후에는 파기합니다.",
          size=7.5, space_before=1)

    doc.save(OUT)
    print("[ok] saved:", OUT)


if __name__ == "__main__":
    main()
